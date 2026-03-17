import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import matplotlib.pyplot as plt
from io import BytesIO

# Configuración
CSV_FILE = "dataset_eventos_50000_v2_sucio.csv"
OUTPUT_DIR = "reportes_generados"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

print(f"Leyendo dataset: {CSV_FILE}...")
df = pd.read_csv(CSV_FILE, encoding="utf-8-sig")

def limpiar_numero(val):
    if pd.isna(val): return 0.0
    val_str = str(val).replace('$', '').replace(' ', '')
    if ',' in val_str and '.' in val_str:
        if val_str.rfind(',') > val_str.rfind('.'):
            val_str = val_str.replace('.', '').replace(',', '.')
        else:
            val_str = val_str.replace(',', '')
    elif ',' in val_str:
        val_str = val_str.replace(',', '.')
    try:
        return float(val_str)
    except:
        return 0.0

def limpiar_texto(val):
    if pd.isna(val): return "Desconocido"
    return str(val).strip().title()

def preparar_datos_limpios(df):
    df['monto_transaccion'] = df['monto_transaccion'].apply(limpiar_numero)
    if 'monto_total' not in df.columns:
        df['monto_total'] = df['monto_transaccion']
    if 'comision' not in df.columns:
        df['comision'] = df['monto_total'] * 0.02
    if 'impuesto' in df.columns:
        df['impuesto'] = df['impuesto'].apply(limpiar_numero)
    if 'ingresos_mensuales' in df.columns:
        df['ingresos_mensuales'] = df['ingresos_mensuales'].apply(limpiar_numero)
    df['es_fraude'] = df['es_fraude'].apply(lambda x: 1 if str(x).strip().lower() in ['1', 'true', 'si', 'yes'] else 0)
    for col in ['canal_transaccion', 'categoria_comercio', 'profesion', 'genero', 'departamento', 'ciudad', 'categoria_tarjeta', 'nombre_comercio', 'estado_civil']:
        if col in df.columns:
            df[col] = df[col].apply(limpiar_texto)
    if 'id_transaccion' not in df.columns:
        df['id_transaccion'] = df.index
    return df

df = preparar_datos_limpios(df)

# --- Preparación de datos (Fechas, Trimestres, Antigüedad) ---
df['fecha_transaccion'] = pd.to_datetime(df['fecha_transaccion'], errors='coerce')
df['fecha_emision'] = pd.to_datetime(df['fecha_emision'], errors='coerce')

df['mes'] = df['fecha_transaccion'].dt.to_period('M')

# Trimestre: Q1, Q2, Q3, Q4
df['trimestre'] = df['fecha_transaccion'].dt.to_period('Q')

# Antigüedad en años (aproximada, basándose en la fecha de emisión de la tarjeta respecto a la fecha de la transacción)
df['antiguedad_anos'] = (df['fecha_transaccion'] - df['fecha_emision']).dt.days / 365.25
df['rango_antiguedad'] = pd.cut(df['antiguedad_anos'], bins=[0, 1, 3, 5, 10, 100], labels=['<1 año', '1-3 años', '3-5 años', '5-10 años', '>10 años'])

# Filtrar fraudes
fraudes = df[df['es_fraude'] == 1]
total_transacciones = len(df)
total_fraudes = len(fraudes)

def format_currency(val):
    if pd.isna(val): return "$0.00"
    return f"${val:,.2f}"

def format_percent(val):
    if pd.isna(val): return "0.00%"
    return f"{val:.2f}%"

def get_chart_stream():
    stream = BytesIO()
    plt.tight_layout()
    plt.savefig(stream, format='png')
    stream.seek(0)
    plt.close()
    return stream

def crear_reporte_kpis():
    global df, fraudes
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading('REPORTE EJECUTIVO DE KPIS Y PREGUNTAS DE NEGOCIO (FRAUDE)', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total de Transacciones Analizadas: {total_transacciones:,}")
    doc.add_paragraph(f"Total de Transacciones Fraudulentas: {total_fraudes:,} ({format_percent((total_fraudes/total_transacciones)*100)})")
    
    # helper de tabla
    def crear_tabla(doc, df_datos, headers):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            
        for index, row in df_datos.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        return table

    # ---------------------------------------------------------
    # Pregunta 1: Fraude por ciudad y variación trimestral
    # ---------------------------------------------------------
    doc.add_heading('1) Fraude por Ciudad y Variación Trimestral', level=1)
    
    # KPI 1 (Tasa de Fraude por Ciudad) y KPI 3 (Monto Refundado)
    ciudad_stats = df.groupby('ciudad').agg(
        total_tx=('id_transaccion', 'count'),
        fraudes=('es_fraude', 'sum')
    )
    ciudad_stats['tasa_fraude'] = (ciudad_stats['fraudes'] / ciudad_stats['total_tx']) * 100
    
    monto_fraude_ciudad = fraudes.groupby('ciudad')['monto_transaccion'].sum()
    
    resumen_ciudad = pd.DataFrame({
        'KPI 1: Tasa Fraude (%)': ciudad_stats['tasa_fraude'].apply(format_percent),
        'KPI 3: Monto Fraudulento': monto_fraude_ciudad.apply(format_currency),
        'Monto_Num': monto_fraude_ciudad.fillna(0)
    }).reset_index().sort_values(by='Monto_Num', ascending=False).head(10)
    
    crear_tabla(doc, resumen_ciudad[['ciudad', 'KPI 1: Tasa Fraude (%)', 'KPI 3: Monto Fraudulento']], ['Ciudad', 'KPI 1 (Tasa)', 'KPI 3 (Monto)'])
    
    plt.figure(figsize=(8, 4))
    plt.bar(resumen_ciudad['ciudad'], resumen_ciudad['Monto_Num'], color='crimson')
    plt.title('Top 10 Ciudades por Monto Fraudulento')
    plt.xlabel('Ciudad')
    plt.ylabel('Monto ($)')
    plt.xticks(rotation=45)
    doc.add_picture(get_chart_stream(), width=Inches(6.0))

    # KPI 2 Variación trimestral (Global para simplificar tabla)
    doc.add_paragraph("KPI 2: Variación Trimestral de Fraude (Global)")
    fraudes_trimestre = fraudes.groupby('trimestre').size().reset_index(name='fraudes')
    fraudes_trimestre['variacion_pct'] = fraudes_trimestre['fraudes'].pct_change() * 100
    
    fraudes_trimestre['Trimestre_str'] = fraudes_trimestre['trimestre'].astype(str)
    fraudes_trimestre['Variación (%)'] = fraudes_trimestre['variacion_pct'].apply(format_percent)
    crear_tabla(doc, fraudes_trimestre[['Trimestre_str', 'fraudes', 'Variación (%)']], ['Trimestre', 'Total Fraudes', 'Variación (%)'])

    plt.figure(figsize=(6, 4))
    plt.plot(fraudes_trimestre['Trimestre_str'], fraudes_trimestre['fraudes'], marker='o', linestyle='-', color='darkred')
    plt.title('Evolución de Fraudes por Trimestre')
    plt.xlabel('Trimestre')
    plt.ylabel('Cantidad de Fraudes')
    plt.grid(True, linestyle='--', alpha=0.7)
    doc.add_picture(get_chart_stream(), width=Inches(5.0))

    # ---------------------------------------------------------
    # Pregunta 2: Tipo de Tarjeta
    # ---------------------------------------------------------
    doc.add_heading('2) Fraudes por Tipo de Tarjeta', level=1)
    
    resumen_tarjeta = pd.DataFrame({
        'Monto_Total': df.groupby('categoria_tarjeta')['monto_transaccion'].sum(),
        'KPI 4: Monto Fraudulento': fraudes.groupby('categoria_tarjeta')['monto_transaccion'].sum(),
        'KPI 6: Número Fraudes': fraudes.groupby('categoria_tarjeta').size()
    }).fillna(0).reset_index()
    resumen_tarjeta.rename(columns={'categoria_tarjeta': 'Tipo'}, inplace=True)
    
    resumen_tarjeta['Monto_Num'] = resumen_tarjeta['KPI 4: Monto Fraudulento']
    resumen_tarjeta['KPI 5: % Fraude (Monto)'] = np.where(resumen_tarjeta['Monto_Total'] > 0, 
                                                          (resumen_tarjeta['Monto_Num'] / resumen_tarjeta['Monto_Total']) * 100, 0)
    
    resumen_tarjeta = resumen_tarjeta.sort_values('Monto_Num', ascending=False)
    
    fmt_tarjeta = resumen_tarjeta.copy()
    fmt_tarjeta['KPI 4: Monto Fraudulento'] = fmt_tarjeta['KPI 4: Monto Fraudulento'].apply(format_currency)
    fmt_tarjeta['KPI 5: % Fraude (Monto)'] = fmt_tarjeta['KPI 5: % Fraude (Monto)'].apply(format_percent)
    
    crear_tabla(doc, fmt_tarjeta[['Tipo', 'KPI 4: Monto Fraudulento', 'KPI 5: % Fraude (Monto)', 'KPI 6: Número Fraudes']], ['Categoría Tarjeta', 'KPI 4 (Monto)', 'KPI 5 (%)', 'KPI 6 (# Fraudes)'])

    if not resumen_tarjeta.empty and sum(resumen_tarjeta['Monto_Num']) > 0:
        plt.figure(figsize=(6, 4))
        plt.pie(resumen_tarjeta['Monto_Num'], labels=resumen_tarjeta['Tipo'], autopct='%1.1f%%', colors=plt.cm.Paired.colors)
        plt.title('Monto Fraudulento por Categoría de Tarjeta')
        doc.add_picture(get_chart_stream(), width=Inches(5.0))

    # ---------------------------------------------------------
    # Pregunta 3: Segmento de Clientes (Profesión + Estado Civil)
    # ---------------------------------------------------------
    doc.add_heading('3) Segmento de Clientes (Profesión + Estado Civil)', level=1)
    
    df['segmento'] = df['profesion'].fillna('') + ' - ' + df['estado_civil'].fillna('')
    fraudes = df[df['es_fraude'] == 1]
    
    resumen_seg = pd.DataFrame({
        'Total_Tx': df.groupby('segmento').size(),
        'KPI 8: # Fraudes': fraudes.groupby('segmento').size(),
        'Monto_Fraude': fraudes.groupby('segmento')['monto_transaccion'].sum()
    }).fillna(0).reset_index()
    
    resumen_seg['Tasa_Num'] = np.where(resumen_seg['Total_Tx'] > 0, (resumen_seg['KPI 8: # Fraudes'] / resumen_seg['Total_Tx']) * 100, 0)
    resumen_seg['KPI 7: Tasa Fraude (%)'] = resumen_seg['Tasa_Num'].apply(format_percent)
    resumen_seg['KPI 9: Promedio Fraude'] = np.where(resumen_seg['KPI 8: # Fraudes'] > 0, resumen_seg['Monto_Fraude'] / resumen_seg['KPI 8: # Fraudes'], 0)
    
    resumen_seg = resumen_seg.sort_values('Tasa_Num', ascending=False).head(10)
    
    fmt_seg = resumen_seg.copy()
    fmt_seg['KPI 9: Promedio Fraude'] = fmt_seg['KPI 9: Promedio Fraude'].apply(format_currency)
    
    crear_tabla(doc, fmt_seg[['segmento', 'KPI 8: # Fraudes', 'KPI 7: Tasa Fraude (%)', 'KPI 9: Promedio Fraude']], ['Segmento', 'KPI 8 (#)', 'KPI 7 (Tasa)', 'KPI 9 (Promedio)'])

    if not resumen_seg.empty:
        plt.figure(figsize=(8, 5))
        plt.barh(resumen_seg['segmento'], resumen_seg['Tasa_Num'], color='teal')
        plt.title('Top 10 Segmentos por Tasa de Fraude (%)')
        plt.xlabel('Tasa de Fraude (%)')
        plt.gca().invert_yaxis()
        doc.add_picture(get_chart_stream(), width=Inches(6.0))

    # ---------------------------------------------------------
    # Pregunta 4: Franja Horaria
    # ---------------------------------------------------------
    doc.add_heading('4) Fraudes por Franja Horaria', level=1)
    doc.add_paragraph(
        "NOTA: El dataset actual (`dataset_eventos_50000.csv`) incluye fechas en formato `YYYY-MM-DD` pero carece de "
        "información sobre horas/minutos en las transacciones. Por tanto, los KPIs 10, 11 y 12 no pueden calcularse "
        "con los datos proporcionados. Se recomienda capturar la franja horaria en futuras extracciones para este análisis."
    ).style = 'Intense Quote'

    # ---------------------------------------------------------
    # Pregunta 5: Promedio y Desviación de Fraude
    # ---------------------------------------------------------
    doc.add_heading('5) Monto Promedio y Desviación Estándar del Fraude', level=1)
    
    promedio_fraude = fraudes['monto_transaccion'].mean() if not fraudes.empty else 0
    std_fraude = fraudes['monto_transaccion'].std() if not fraudes.empty else 0
    
    doc.add_paragraph(f"KPI 13 (Monto Promedio Fraudulento): {format_currency(promedio_fraude)}")
    doc.add_paragraph(f"KPI 14 (Desviación Estándar): {format_currency(std_fraude)} (Muestra la alta variabilidad en los montos de fraude).")

    # ---------------------------------------------------------
    # Pregunta 6: Comercios con Mayor Fraude
    # ---------------------------------------------------------
    doc.add_heading('6) Comercios con Mayor Fraude', level=1)
    
    resumen_com = pd.DataFrame({
        'Total_Tx': df.groupby('nombre_comercio').size(),
        'KPI 15: # Fraudes': fraudes.groupby('nombre_comercio').size(),
        'KPI 16: Monto': fraudes.groupby('nombre_comercio')['monto_transaccion'].sum()
    }).fillna(0).reset_index()
    
    resumen_com['Num_Fraudes'] = resumen_com['KPI 15: # Fraudes']
    resumen_com['KPI 17: Tasa Fraude (%)'] = np.where(resumen_com['Total_Tx'] > 0, (resumen_com['Num_Fraudes'] / resumen_com['Total_Tx']) * 100, 0)
    
    resumen_com = resumen_com[resumen_com['Num_Fraudes'] > 0].sort_values('Num_Fraudes', ascending=False).head(5)
    
    fmt_com = resumen_com.copy()
    fmt_com['KPI 16: Monto'] = fmt_com['KPI 16: Monto'].apply(format_currency)
    fmt_com['KPI 17: Tasa Fraude (%)'] = fmt_com['KPI 17: Tasa Fraude (%)'].apply(format_percent)
    
    crear_tabla(doc, fmt_com[['nombre_comercio', 'KPI 15: # Fraudes', 'KPI 16: Monto', 'KPI 17: Tasa Fraude (%)']], ['Comercio', 'KPI 15 (#)', 'KPI 16 (Monto)', 'KPI 17 (Tasa)'])

    if not resumen_com.empty:
        plt.figure(figsize=(7, 4))
        plt.bar(resumen_com['nombre_comercio'], resumen_com['Num_Fraudes'], color='coral')
        plt.title('Top 5 Comercios con Mayor Número de Fraudes')
        plt.ylabel('Cantidad de Fraudes')
        plt.xticks(rotation=45)
        doc.add_picture(get_chart_stream(), width=Inches(5.5))

    # ---------------------------------------------------------
    # Pregunta 7: Mes con Mayor Monto Fraudulento
    # ---------------------------------------------------------
    doc.add_heading('7) Mes con Mayor Monto de Fraude', level=1)
    
    resumen_mes = pd.DataFrame({
        'KPI 19: # Fraudes': fraudes.groupby('mes').size(),
        'KPI 18: Monto': fraudes.groupby('mes')['monto_transaccion'].sum()
    }).fillna(0).reset_index()
    
    resumen_mes['Monto_Num'] = resumen_mes['KPI 18: Monto']
    resumen_mes['mes_str'] = resumen_mes['mes'].astype(str)
    resumen_mes['KPI 20: Variación (%)'] = resumen_mes['Monto_Num'].pct_change() * 100
    
    fmt_mes = resumen_mes.copy().sort_values('Monto_Num', ascending=False).head(6)
    fmt_mes['KPI 18: Monto'] = fmt_mes['KPI 18: Monto'].apply(format_currency)
    fmt_mes['KPI 20: Variación (%)'] = fmt_mes['KPI 20: Variación (%)'].apply(format_percent)
    
    crear_tabla(doc, fmt_mes[['mes_str', 'KPI 19: # Fraudes', 'KPI 18: Monto', 'KPI 20: Variación (%)']], ['Mes', 'KPI 19 (#)', 'KPI 18 (Monto)', 'KPI 20 (Var %)'])

    if not resumen_mes.empty:
        resumen_mes_crono = resumen_mes.sort_values('mes_str')
        plt.figure(figsize=(7, 4))
        plt.plot(resumen_mes_crono['mes_str'], resumen_mes_crono['Monto_Num'], marker='s', color='purple')
        plt.title('Monto de Fraude por Mes')
        plt.xlabel('Mes')
        plt.ylabel('Monto ($)')
        plt.xticks(rotation=45)
        plt.grid(True, linestyle=':', alpha=0.6)
        doc.add_picture(get_chart_stream(), width=Inches(5.5))

    # ---------------------------------------------------------
    # Pregunta 8: Fraude según Antigüedad
    # ---------------------------------------------------------
    doc.add_heading('8) Tasa de Fraude según Antigüedad del Cliente', level=1)
    
    resumen_antig = pd.DataFrame({
        'Total_Tx': df.groupby('rango_antiguedad', observed=True).size(),
        'Fraudes': fraudes.groupby('rango_antiguedad', observed=True).size(),
        'Monto': fraudes.groupby('rango_antiguedad', observed=True)['monto_transaccion'].sum(),
        'Clientes': fraudes.groupby('rango_antiguedad', observed=True)['nombre_cliente'].nunique()
    }).fillna(0).reset_index()
    
    resumen_antig['Tasa_Num'] = np.where(resumen_antig['Total_Tx'] > 0, (resumen_antig['Fraudes'] / resumen_antig['Total_Tx']) * 100, 0)
    
    fmt_antig = resumen_antig.copy()
    fmt_antig['KPI 21: Tasa (%)'] = fmt_antig['Tasa_Num'].apply(format_percent)
    fmt_antig['KPI 22: Monto'] = fmt_antig['Monto'].apply(format_currency)
    
    crear_tabla(doc, fmt_antig[['rango_antiguedad', 'KPI 21: Tasa (%)', 'KPI 22: Monto', 'Clientes']], ['Antigüedad', 'KPI 21 (Tasa)', 'KPI 22 (Monto)', 'KPI 23 (Clientes)'])

    if not resumen_antig.empty:
        plt.figure(figsize=(6, 4))
        plt.bar(resumen_antig['rango_antiguedad'].astype(str), resumen_antig['Tasa_Num'], color='dodgerblue')
        plt.title('Tasa de Fraude por Antigüedad del Cliente (%)')
        plt.xlabel('Antigüedad')
        plt.ylabel('Tasa de Fraude (%)')
        plt.xticks(rotation=15)
        doc.add_picture(get_chart_stream(), width=Inches(5.0))

    # ---------------------------------------------------------
    # Pregunta 9: Ranking Ciudades y Evolución (Resumen)
    # ---------------------------------------------------------
    doc.add_heading('9) Top Ciudades con Mayor Fraude Mensual', level=1)
    doc.add_paragraph("Nota: KPI 24, 25 y 26 consolidados en el Top 3:")
    
    monto_fraude_ciudad = fraudes.groupby('ciudad')['monto_transaccion'].sum()
    top_ciudades = monto_fraude_ciudad.sort_values(ascending=False).head(3).index
    
    fraude_ciudades = fraudes[fraudes['ciudad'].isin(top_ciudades)]
    if not fraude_ciudades.empty:
        serie_temporal = fraude_ciudades.groupby(['ciudad', 'trimestre'])['monto_transaccion'].sum().reset_index()
        serie_temporal['trimestre_str'] = serie_temporal['trimestre'].astype(str)
        
        fmt_serie = serie_temporal.copy()
        fmt_serie['Monto (KPI 24)'] = fmt_serie['monto_transaccion'].apply(format_currency)
        
        crear_tabla(doc, fmt_serie[['ciudad', 'trimestre_str', 'Monto (KPI 24)']], ['Ciudad', 'Trimestre', 'Monto Fraude'])

        plt.figure(figsize=(8, 4))
        for ciudad in top_ciudades:
            datos_ciudad = serie_temporal[serie_temporal['ciudad'] == ciudad]
            if not datos_ciudad.empty:
                plt.plot(datos_ciudad['trimestre_str'], datos_ciudad['monto_transaccion'], marker='o', label=ciudad)
        plt.title('Evolución de Fraude Mensual para Top 3 Ciudades')
        plt.xlabel('Trimestre')
        plt.ylabel('Monto ($)')
        plt.legend()
        plt.grid(True, linestyle='--', alpha=0.5)
        doc.add_picture(get_chart_stream(), width=Inches(6.0))

    # Save
    reporte_path = os.path.join(OUTPUT_DIR, 'Reporte_KPIs_Negocio.docx')
    try:
        doc.save(reporte_path)
    except PermissionError:
        import time
        reporte_path = os.path.join(OUTPUT_DIR, f'Reporte_KPIs_Negocio_{int(time.time())}.docx')
        doc.save(reporte_path)
    print(f"\nReporte Ejecutivo generado exitosamente en: '{reporte_path}'")

if __name__ == "__main__":
    import warnings
    warnings.filterwarnings('ignore') # Omitir warnings de pandas para cortes de fechas vacias
    crear_reporte_kpis()
