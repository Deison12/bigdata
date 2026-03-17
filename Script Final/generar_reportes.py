import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import matplotlib.pyplot as plt
from io import BytesIO

# Configuración
import pathlib
BASE_DIR = pathlib.Path(__file__).resolve().parent
CSV_FILE = str(BASE_DIR / "dataset_eventos_50000_v2_sucio.csv")
OUTPUT_DIR = str(BASE_DIR / "reportes_generados")

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

print(f"Leyendo dataset: {CSV_FILE}...")
df = pd.read_csv(CSV_FILE, encoding="utf-8-sig")

def limpiar_numero(val):
    if pd.isna(val): return 0.0
    val_str = str(val).replace('$', '').replace(' ', '')
    if ',' in val_str and '.' in val_str:
        if val_str.rfind(',') > val_str.rfind('.'):
            val_str = val_str.replace('.', '').replace(',', '.')
        else:
            val_str = val_str.replace(',', '')
    elif ',' in val_str:
        val_str = val_str.replace(',', '.')
    try:
        return float(val_str)
    except:
        return 0.0

def limpiar_texto(val):
    if pd.isna(val): return "Desconocido"
    return str(val).strip().title()

def preparar_datos_limpios(df):
    df['monto_transaccion'] = df['monto_transaccion'].apply(limpiar_numero)
    if 'monto_total' not in df.columns:
        df['monto_total'] = df['monto_transaccion']
    if 'comision' not in df.columns:
        df['comision'] = df['monto_total'] * 0.02
    if 'impuesto' in df.columns:
        df['impuesto'] = df['impuesto'].apply(limpiar_numero)
    if 'ingresos_mensuales' in df.columns:
        df['ingresos_mensuales'] = df['ingresos_mensuales'].apply(limpiar_numero)
    df['es_fraude'] = df['es_fraude'].apply(lambda x: 1 if str(x).strip().lower() in ['1', 'true', 'si', 'yes'] else 0)
    for col in ['canal_transaccion', 'categoria_comercio', 'profesion', 'genero', 'departamento', 'ciudad', 'categoria_tarjeta', 'nombre_comercio', 'estado_civil']:
        if col in df.columns:
            df[col] = df[col].apply(limpiar_texto)
    if 'id_transaccion' not in df.columns:
        df['id_transaccion'] = df.index
    return df

df = preparar_datos_limpios(df)

def get_chart_stream():
    stream = BytesIO()
    plt.tight_layout()
    plt.savefig(stream, format='png')
    stream.seek(0)
    plt.close()
    return stream

def guardar_reporte(doc, output_dir, filename):
    import time
    path = os.path.join(output_dir, filename)
    try:
        doc.save(path)
    except PermissionError:
        path = os.path.join(output_dir, filename.replace('.docx', f'_{int(time.time())}.docx'))
        doc.save(path)
    return path

# 1. Reporte de Fraudes
def generar_reporte_fraudes():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: DETECCIÓN Y PREVENCIÓN DE FRAUDES', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    fraudes = df[df['es_fraude'] == 1]
    total_fraudes = len(fraudes)
    tasa_fraude = (total_fraudes / len(df)) * 100
    
    doc.add_paragraph(
        f"Durante el periodo analizado, se han identificado {total_fraudes} transacciones con indicadores de fraude, "
        f"lo que representa una tasa del {tasa_fraude:.2f}% sobre el volumen total de {len(df):,} operaciones. "
        "Este reporte desglosa los canales y categorías más vulnerables para la toma de decisiones estratégicas en seguridad."
    )
    
    # Análisis por Canal
    doc.add_heading('2. Análisis por Canal de Transacción', level=1)
    canal_fraude = fraudes['canal_transaccion'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Canal'
    hdr_cells[1].text = 'Número de Fraudes'
    
    for canal, count in canal_fraude.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(canal)
        row_cells[1].text = f"{count:,}"
    
    plt.figure(figsize=(6, 4))
    canal_fraude.plot(kind='bar', color='salmon')
    plt.title('Número de Fraudes por Canal')
    plt.xlabel('Canal')
    plt.ylabel('Cantidad')
    plt.xticks(rotation=45)
    doc.add_picture(get_chart_stream(), width=Inches(5.5))

    # Análisis por Categoría
    doc.add_heading('3. Análisis por Categoría de Comercio', level=1)
    cat_fraude = fraudes['categoria_comercio'].value_counts()
    
    table_cat = doc.add_table(rows=1, cols=2)
    table_cat.style = 'Light Shading Accent 1'
    hdr_cells = table_cat.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Incidencia de Fraude'
    
    for cat, count in cat_fraude.items():
        row_cells = table_cat.add_row().cells
        row_cells[0].text = str(cat)
        row_cells[1].text = f"{count:,}"

    plt.figure(figsize=(6, 4))
    cat_fraude.plot(kind='pie', autopct='%1.1f%%', colormap='Pastel1')
    plt.title('Incidencia de Fraude por Categoría')
    plt.ylabel('')
    doc.add_picture(get_chart_stream(), width=Inches(5.0))

    # Recomendaciones
    doc.add_heading('4. Conclusiones y Recomendaciones', level=1)
    doc.add_paragraph("1. Reforzar los protocolos de validación en los canales con mayor incidencia.")
    doc.add_paragraph("2. Implementar reglas de monitoreo en tiempo real para las categorías de comercio críticas.")
    
    guardar_reporte(doc, OUTPUT_DIR, 'Reporte_Fraudes_Ejecutivo.docx')
    print("Reporte de Fraudes Ejecutivo generado.")

# 2. Reporte Demográfico
def generar_reporte_demografico():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: PERFIL DEMOGRÁFICO DE CLIENTES', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        "Este informe presenta una visión consolidada de la base de clientes, analizando la segmentación por "
        "profesión y el poder adquisitivo promedio. Estos datos permiten alinear las campañas de marketing "
        "y la oferta de productos financieros."
    )
    
    # Distribución Profesional
    doc.add_heading('2. Distribución por Segmento Profesional', level=1)
    profesiones = df['profesion'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Profesión'
    hdr_cells[1].text = 'Cantidad de Clientes'
    
    for prof, count in profesiones.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(prof)
        row_cells[1].text = f"{count:,}"

    plt.figure(figsize=(7, 4))
    profesiones.plot(kind='barh', color='skyblue')
    plt.title('Cantidad de Clientes por Profesión')
    plt.xlabel('Cantidad')
    plt.ylabel('Profesión')
    doc.add_picture(get_chart_stream(), width=Inches(6.0))
    
    # Nivel Económico
    doc.add_heading('3. Ingreso Mensual Promedio por Género', level=1)
    ingresos = df.groupby('genero')['ingresos_mensuales'].mean()
    
    table_ing = doc.add_table(rows=1, cols=2)
    table_ing.style = 'Table Grid'
    hdr_cells = table_ing.rows[0].cells
    hdr_cells[0].text = 'Género'
    hdr_cells[1].text = 'Ingreso Promedio'
    
    for gen, val in ingresos.items():
        row_cells = table_ing.add_row().cells
        row_cells[0].text = str(gen)
        row_cells[1].text = f"${val:,.2f}"

    plt.figure(figsize=(5, 4))
    ingresos.plot(kind='bar', color='lightcoral')
    plt.title('Ingreso Mensual Promedio por Género')
    plt.ylabel('Ingreso ($)')
    plt.xticks(rotation=0)
    doc.add_picture(get_chart_stream(), width=Inches(4.5))
        
    guardar_reporte(doc, OUTPUT_DIR, 'Reporte_Demografico_Ejecutivo.docx')
    print("Reporte Demográfico Ejecutivo generado.")

# 3. Reporte Operativo
def generar_reporte_operativo():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: ANÁLISIS OPERATIVO TERRITORIAL', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        "Se presenta el análisis de la distribución geográfica de las operaciones. Identificar los focos de mayor "
        "actividad permite optimizar la logística de servicios y la presencia comercial en las regiones clave."
    )
    
    # Actividad por Departamento
    doc.add_heading('2. Actividad por Departamento', level=1)
    deptos = df['departamento'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 2'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Departamento'
    hdr_cells[1].text = 'Nivel de Operaciones'
    
    for depto, count in deptos.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(depto)
        row_cells[1].text = f"{count:,} tx"

    plt.figure(figsize=(7, 4))
    deptos.head(10).plot(kind='bar', color='mediumseagreen')
    plt.title('Top 10 Departamentos por Operaciones')
    plt.ylabel('Cantidad de Transacciones')
    plt.xticks(rotation=45)
    doc.add_picture(get_chart_stream(), width=Inches(6.0))
        
    # Top Ciudades
    doc.add_heading('3. Ciudades con Mayor Volumen de Transacción', level=1)
    ciudades = df['ciudad'].value_counts().head(5)
    
    for ciudad, count in ciudades.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{ciudad}: ").bold = True
        p.add_run(f"{count:,} transacciones")

    plt.figure(figsize=(6, 4))
    ciudades.plot(kind='pie', autopct='%1.1f%%', colormap='Set3')
    plt.title('Top 5 Ciudades por Volumen de Transacción')
    plt.ylabel('')
    doc.add_picture(get_chart_stream(), width=Inches(5.0))
        
    guardar_reporte(doc, OUTPUT_DIR, 'Reporte_Operativo_Ejecutivo.docx')
    print("Reporte Operativo Ejecutivo generado.")

# 4. Reporte Financiero
def generar_reporte_financiero():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: DESEMPEÑO FINANCIERO Y RENTABILIDAD', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Estados Financieros Consolidados', level=1)
    total_monto = df['monto_total'].sum()
    total_comision = df['comision'].sum()
    total_impuesto = df['impuesto'].sum()
    
    table_res = doc.add_table(rows=3, cols=2)
    table_res.style = 'Colorful List Accent 3'
    
    table_res.rows[0].cells[0].text = "Volumen Total Transado"
    table_res.rows[0].cells[1].text = f"${total_monto:,.2f}"
    
    table_res.rows[1].cells[0].text = "Total Comisiones"
    table_res.rows[1].cells[1].text = f"${total_comision:,.2f}"
    
    table_res.rows[2].cells[0].text = "Total Impuestos Recaudados"
    table_res.rows[2].cells[1].text = f"${total_impuesto:,.2f}"

    plt.figure(figsize=(6, 4))
    pd.Series({'Comisiones': total_comision, 'Impuestos': total_impuesto}).plot(kind='bar', color=['gold', 'orange'])
    plt.title('Total Comisiones vs Impuestos')
    plt.ylabel('Monto ($)')
    plt.xticks(rotation=0)
    doc.add_picture(get_chart_stream(), width=Inches(5.0))
    
    # Rentabilidad por Producto
    doc.add_heading('2. Carga Financiera por Categoría de Tarjeta', level=1)
    cat_tarjeta = df.groupby('categoria_tarjeta')['monto_total'].sum()
    
    table_tar = doc.add_table(rows=1, cols=2)
    table_tar.style = 'Table Grid'
    hdr_cells = table_tar.rows[0].cells
    hdr_cells[0].text = 'Tipo de Tarjeta'
    hdr_cells[1].text = 'Monto Acumulado'
    
    for cat, val in cat_tarjeta.items():
        row_cells = table_tar.add_row().cells
        row_cells[0].text = str(cat)
        row_cells[1].text = f"${val:,.2f}"

    plt.figure(figsize=(6, 4))
    cat_tarjeta.plot(kind='bar', color='purple')
    plt.title('Carga Financiera por Categoría de Tarjeta')
    plt.ylabel('Monto Acumulado ($)')
    plt.xlabel('Categoría de Tarjeta')
    plt.xticks(rotation=45)
    doc.add_picture(get_chart_stream(), width=Inches(5.5))
        
    guardar_reporte(doc, OUTPUT_DIR, 'Reporte_Financiero_Ejecutivo.docx')
    print("Reporte Financiero Ejecutivo generado.")

if __name__ == "__main__":
    generar_reporte_fraudes()
    generar_reporte_demografico()
    generar_reporte_operativo()
    generar_reporte_financiero()
    print(f"\nProceso finalizado. Reportes guardados en '{OUTPUT_DIR}'")
