import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Configuración
CSV_FILE = "dataset_eventos_50000.csv"
OUTPUT_DIR = "reportes_generados"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

print(f"Leyendo dataset: {CSV_FILE}...")
df = pd.read_csv(CSV_FILE, encoding="utf-8-sig")

# --- Preparación de datos (Fechas, Trimestres, Antigüedad) ---
df['fecha_transaccion'] = pd.to_datetime(df['fecha_transaccion'])
df['fecha_emision'] = pd.to_datetime(df['fecha_emision'])

df['mes'] = df['fecha_transaccion'].dt.to_period('M')

# Trimestre: Q1, Q2, Q3, Q4
df['trimestre'] = df['fecha_transaccion'].dt.to_period('Q')

# Antigüedad en años (aproximada, basándose en la fecha de emisión de la tarjeta respecto a la fecha de la transacción)
df['antiguedad_anos'] = (df['fecha_transaccion'] - df['fecha_emision']).dt.days / 365.25
df['rango_antiguedad'] = pd.cut(df['antiguedad_anos'], bins=[0, 1, 3, 5, 10, 100], labels=['<1 año', '1-3 años', '3-5 años', '5-10 años', '>10 años'])

# Filtrar fraudes
fraudes = df[df['es_fraude'] == 1]
total_transacciones = len(df)
total_fraudes = len(fraudes)

def format_currency(val):
    return f"${val:,.2f}"

def format_percent(val):
    return f"{val:.2f}%"

def crear_reporte_kpis():
    doc = Document()
    
    # Título Principal
    titulo = doc.add_heading('REPORTE EJECUTIVO DE KPIS Y PREGUNTAS DE NEGOCIO (FRAUDE)', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total de Transacciones Analizadas: {total_transacciones:,}")
    doc.add_paragraph(f"Total de Transacciones Fraudulentas: {total_fraudes:,} ({format_percent((total_fraudes/total_transacciones)*100)})")
    
    # helper de tabla
    def crear_tabla(doc, df_datos, headers):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            
        for index, row in df_datos.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)
        return table

    # ---------------------------------------------------------
    # Pregunta 1: Fraude por ciudad y variación trimestral
    # ---------------------------------------------------------
    doc.add_heading('1) Fraude por Ciudad y Variación Trimestral', level=1)
    
    # KPI 1 (Tasa de Fraude por Ciudad) y KPI 3 (Monto Refundado)
    ciudad_stats = df.groupby('ciudad').agg(
        total_tx=('id_transaccion', 'count'),
        fraudes=('es_fraude', 'sum')
    )
    ciudad_stats['tasa_fraude'] = (ciudad_stats['fraudes'] / ciudad_stats['total_tx']) * 100
    
    monto_fraude_ciudad = fraudes.groupby('ciudad')['monto_transaccion'].sum()
    
    resumen_ciudad = pd.DataFrame({
        'KPI 1: Tasa Fraude (%)': ciudad_stats['tasa_fraude'].apply(lambda x: format_percent(x)),
        'KPI 3: Monto Fraudulento': monto_fraude_ciudad.apply(lambda x: format_currency(x) if pd.notnull(x) else "$0.00")
    }).reset_index().fillna("$0.00").sort_values(by='KPI 3: Monto Fraudulento', ascending=False).head(10)
    
    crear_tabla(doc, resumen_ciudad, ['Ciudad', 'KPI 1 (Tasa)', 'KPI 3 (Monto)'])
    
    # KPI 2 Variación trimestral (Global para simplificar tabla)
    doc.add_paragraph("KPI 2: Variación Trimestral de Fraude (Global)")
    fraudes_trimestre = fraudes.groupby('trimestre').size().reset_index(name='fraudes')
    fraudes_trimestre['variacion_pct'] = fraudes_trimestre['fraudes'].pct_change() * 100
    
    fraudes_trimestre['Trimestre'] = fraudes_trimestre['trimestre'].astype(str)
    fraudes_trimestre['Variación (%)'] = fraudes_trimestre['variacion_pct'].apply(lambda x: format_percent(x) if pd.notnull(x) else "N/A")
    crear_tabla(doc, fraudes_trimestre[['Trimestre', 'fraudes', 'Variación (%)']], ['Trimestre', 'Total Fraudes', 'Variación (%)'])


    # ---------------------------------------------------------
    # Pregunta 2: Tipo de Tarjeta
    # ---------------------------------------------------------
    doc.add_heading('2) Fraudes por Tipo de Tarjeta', level=1)
    
    monto_total_tarjeta = df.groupby('categoria_tarjeta')['monto_transaccion'].sum()
    monto_fraude_tarjeta = fraudes.groupby('categoria_tarjeta')['monto_transaccion'].sum()
    conteo_fraude_tarjeta = fraudes.groupby('categoria_tarjeta').size()
    
    resumen_tarjeta = pd.DataFrame({
        'Tipo': monto_total_tarjeta.index,
        'KPI 4: Monto Fraudulento': monto_fraude_tarjeta.apply(lambda x: format_currency(x)).values,
        'KPI 5: % Fraude (Monto)': ((monto_fraude_tarjeta / monto_total_tarjeta) * 100).apply(lambda x: format_percent(x)).values,
        'KPI 6: Número Fraudes': conteo_fraude_tarjeta.values
    }).sort_values('KPI 4: Monto Fraudulento', ascending=False)
    
    crear_tabla(doc, resumen_tarjeta, ['Categoría Tarjeta', 'KPI 4 (Monto)', 'KPI 5 (%)', 'KPI 6 (# Fraudes)'])

    # ---------------------------------------------------------
    # Pregunta 3: Segmento de Clientes (Profesión + Estado Civil)
    # ---------------------------------------------------------
    doc.add_heading('3) Segmento de Clientes (Profesión + Estado Civil)', level=1)
    
    df['segmento'] = df['profesion'] + ' - ' + df['estado_civil']
    fraudes['segmento'] = fraudes['profesion'] + ' - ' + fraudes['estado_civil']
    
    seg_total = df.groupby('segmento').size()
    seg_fraude = fraudes.groupby('segmento').size()
    monto_seg_fraude = fraudes.groupby('segmento')['monto_transaccion'].sum()
    
    resumen_seg = pd.DataFrame({
        'KPI 8: # Fraudes': seg_fraude,
        'KPI 7: Tasa Fraude (%)': (seg_fraude / seg_total) * 100,
        'KPI 9: Promedio Fraude': monto_seg_fraude / seg_fraude
    }).reset_index().dropna().sort_values('KPI 7: Tasa Fraude (%)', ascending=False).head(10)
    
    resumen_seg['KPI 7: Tasa Fraude (%)'] = resumen_seg['KPI 7: Tasa Fraude (%)'].apply(format_percent)
    resumen_seg['KPI 9: Promedio Fraude'] = resumen_seg['KPI 9: Promedio Fraude'].apply(format_currency)
    
    crear_tabla(doc, resumen_seg, ['Segmento', 'KPI 8 (#)', 'KPI 7 (Tasa)', 'KPI 9 (Promedio)'])

    # ---------------------------------------------------------
    # Pregunta 4: Franja Horaria
    # ---------------------------------------------------------
    doc.add_heading('4) Fraudes por Franja Horaria', level=1)
    doc.add_paragraph(
        "NOTA: El dataset actual (`dataset_eventos_50000.csv`) incluye fechas en formato `YYYY-MM-DD` pero carece de "
        "información sobre horas/minutos en las transacciones. Por tanto, los KPIs 10, 11 y 12 no pueden calcularse "
        "con los datos proporcionados. Se recomienda capturar la franja horaria en futuras extracciones para este análisis."
    ).style = 'Intense Quote'

    # ---------------------------------------------------------
    # Pregunta 5: Promedio y Desviación de Fraude
    # ---------------------------------------------------------
    doc.add_heading('5) Monto Promedio y Desviación Estándar del Fraude', level=1)
    
    # KPI 13 (Monto Promedio Fraudulento) y KPI 14 (Desviación estándar)
    promedio_fraude = fraudes['monto_transaccion'].mean()
    std_fraude = fraudes['monto_transaccion'].std()
    
    doc.add_paragraph(f"KPI 13 (Monto Promedio Fraudulento): {format_currency(promedio_fraude)}")
    doc.add_paragraph(f"KPI 14 (Desviación Estándar): {format_currency(std_fraude)} (Muestra la alta variabilidad en los montos de fraude).")

    # ---------------------------------------------------------
    # Pregunta 6: Comercios con Mayor Fraude
    # ---------------------------------------------------------
    doc.add_heading('6) Comercios con Mayor Fraude', level=1)
    
    com_total = df.groupby('nombre_comercio').size()
    com_fraude = fraudes.groupby('nombre_comercio').size()
    com_monto = fraudes.groupby('nombre_comercio')['monto_transaccion'].sum()
    
    resumen_com = pd.DataFrame({
        'KPI 15: # Fraudes': com_fraude,
        'KPI 16: Monto': com_monto,
        'KPI 17: Tasa Fraude (%)': (com_fraude / com_total) * 100
    }).reset_index().dropna().sort_values('KPI 15: # Fraudes', ascending=False).head(5)
    
    resumen_com['KPI 16: Monto'] = resumen_com['KPI 16: Monto'].apply(format_currency)
    resumen_com['KPI 17: Tasa Fraude (%)'] = resumen_com['KPI 17: Tasa Fraude (%)'].apply(format_percent)
    
    crear_tabla(doc, resumen_com, ['Comercio', 'KPI 15 (#)', 'KPI 16 (Monto)', 'KPI 17 (Tasa)'])

    # ---------------------------------------------------------
    # Pregunta 7: Mes con Mayor Monto Fraudulento
    # ---------------------------------------------------------
    doc.add_heading('7) Mes con Mayor Monto de Fraude', level=1)
    
    monto_mes = fraudes.groupby('mes')['monto_transaccion'].sum()
    conteo_mes = fraudes.groupby('mes').size()
    
    resumen_mes = pd.DataFrame({
        'KPI 19: # Fraudes': conteo_mes,
        'KPI 18: Monto': monto_mes
    }).reset_index()
    
    resumen_mes['KPI 20: Variación (%)'] = resumen_mes['KPI 18: Monto'].pct_change() * 100
    
    resumen_mes['mes'] = resumen_mes['mes'].astype(str)
    resumen_mes['KPI 18: Monto'] = resumen_mes['KPI 18: Monto'].apply(format_currency)
    resumen_mes['KPI 20: Variación (%)'] = resumen_mes['KPI 20: Variación (%)'].apply(lambda x: format_percent(x) if pd.notnull(x) else "N/A")
    
    # Ordenamos descendente para encontrar el mayor más fácil pero mostramos cronológicamente si gustan. (Mostrando top 5 montos)
    resumen_mes = resumen_mes.sort_values('KPI 18: Monto', ascending=False).head(6) 
    
    crear_tabla(doc, resumen_mes, ['Mes', 'KPI 19 (#)', 'KPI 18 (Monto)', 'KPI 20 (Var %)'])

    # ---------------------------------------------------------
    # Pregunta 8: Fraude según Antigüedad
    # ---------------------------------------------------------
    doc.add_heading('8) Tasa de Fraude según Antigüedad del Cliente', level=1)
    
    antig_total = df.groupby('rango_antiguedad', observed=True).size()
    antig_fraude = fraudes.groupby('rango_antiguedad', observed=True).size()
    monto_antig = fraudes.groupby('rango_antiguedad', observed=True)['monto_transaccion'].sum()
    clientes_antig = fraudes.groupby('rango_antiguedad', observed=True)['nombre_cliente'].nunique()
    
    resumen_antig = pd.DataFrame({
        'Rango': antig_total.index,
        'KPI 21: Tasa (%)': ((antig_fraude / antig_total) * 100).values,
        'KPI 22: Monto': monto_antig.values,
        'KPI 23: Clientes Únicos': clientes_antig.values
    })
    
    resumen_antig['KPI 21: Tasa (%)'] = resumen_antig['KPI 21: Tasa (%)'].apply(format_percent)
    resumen_antig['KPI 22: Monto'] = resumen_antig['KPI 22: Monto'].apply(format_currency)
    
    crear_tabla(doc, resumen_antig, ['Antigüedad', 'KPI 21 (Tasa)', 'KPI 22 (Monto)', 'KPI 23 (Clientes)'])

    # ---------------------------------------------------------
    # Pregunta 9: Ranking Ciudades y Evolución (Resumen)
    # ---------------------------------------------------------
    doc.add_heading('9) Top Ciudades con Mayor Fraude Mensual', level=1)
    doc.add_paragraph("Nota: KPI 24, 25 y 26 consolidados en el Top 3:")
    
    top_ciudades = monto_fraude_ciudad.sort_values(ascending=False).head(3).index
    
    fraude_ciudades = fraudes[fraudes['ciudad'].isin(top_ciudades)]
    serie_temporal = fraude_ciudades.groupby(['ciudad', 'trimestre'])['monto_transaccion'].sum().reset_index()
    
    serie_temporal['trimestre'] = serie_temporal['trimestre'].astype(str)
    serie_temporal['Monto (KPI 24)'] = serie_temporal['monto_transaccion'].apply(format_currency)
    
    crear_tabla(doc, serie_temporal[['ciudad', 'trimestre', 'Monto (KPI 24)']], ['Ciudad', 'Trimestre', 'Monto Fraude'])

    # Save
    reporte_path = os.path.join(OUTPUT_DIR, 'Reporte_KPIs_Negocio.docx')
    doc.save(reporte_path)
    print(f"\nReporte Ejecutivo generado exitosamente en: '{reporte_path}'")

if __name__ == "__main__":
    import warnings
    warnings.filterwarnings('ignore') # Omitir warnings de pandas para cortes de fechas vacias
    crear_reporte_kpis()
