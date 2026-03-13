import pandas as pd
from docx import Document
from docx.shared import Inches
import os

# Configuración
CSV_FILE = "dataset_eventos_50000.csv"
OUTPUT_DIR = "reportes_generados"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

print(f"Leyendo dataset: {CSV_FILE}...")
df = pd.read_csv(CSV_FILE, encoding="utf-8-sig")

# 1. Reporte de Fraudes
def generar_reporte_fraudes():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: DETECCIÓN Y PREVENCIÓN DE FRAUDES', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    fraudes = df[df['es_fraude'] == 1]
    total_fraudes = len(fraudes)
    tasa_fraude = (total_fraudes / len(df)) * 100
    
    doc.add_paragraph(
        f"Durante el periodo analizado, se han identificado {total_fraudes} transacciones con indicadores de fraude, "
        f"lo que representa una tasa del {tasa_fraude:.2f}% sobre el volumen total de {len(df):,} operaciones. "
        "Este reporte desglosa los canales y categorías más vulnerables para la toma de decisiones estratégicas en seguridad."
    )
    
    # Análisis por Canal
    doc.add_heading('2. Análisis por Canal de Transacción', level=1)
    canal_fraude = fraudes['canal_transaccion'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Canal'
    hdr_cells[1].text = 'Número de Fraudes'
    
    for canal, count in canal_fraude.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(canal)
        row_cells[1].text = f"{count:,}"

    # Análisis por Categoría
    doc.add_heading('3. Análisis por Categoría de Comercio', level=1)
    cat_fraude = fraudes['categoria_comercio'].value_counts()
    
    table_cat = doc.add_table(rows=1, cols=2)
    table_cat.style = 'Light Shading Accent 1'
    hdr_cells = table_cat.rows[0].cells
    hdr_cells[0].text = 'Categoría'
    hdr_cells[1].text = 'Incidencia de Fraude'
    
    for cat, count in cat_fraude.items():
        row_cells = table_cat.add_row().cells
        row_cells[0].text = str(cat)
        row_cells[1].text = f"{count:,}"

    # Recomendaciones
    doc.add_heading('4. Conclusiones y Recomendaciones', level=1)
    doc.add_paragraph("1. Reforzar los protocolos de validación en los canales con mayor incidencia.")
    doc.add_paragraph("2. Implementar reglas de monitoreo en tiempo real para las categorías de comercio críticas.")
    
    doc.save(os.path.join(OUTPUT_DIR, 'Reporte_Fraudes_Ejecutivo.docx'))
    print("Reporte de Fraudes Ejecutivo generado.")

# 2. Reporte Demográfico
def generar_reporte_demografico():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: PERFIL DEMOGRÁFICO DE CLIENTES', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        "Este informe presenta una visión consolidada de la base de clientes, analizando la segmentación por "
        "profesión y el poder adquisitivo promedio. Estos datos permiten alinear las campañas de marketing "
        "y la oferta de productos financieros."
    )
    
    # Distribución Profesional
    doc.add_heading('2. Distribución por Segmento Profesional', level=1)
    profesiones = df['profesion'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Profesión'
    hdr_cells[1].text = 'Cantidad de Clientes'
    
    for prof, count in profesiones.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(prof)
        row_cells[1].text = f"{count:,}"
    
    # Nivel Económico
    doc.add_heading('3. Ingreso Mensual Promedio por Género', level=1)
    ingresos = df.groupby('genero')['ingresos_mensuales'].mean()
    
    table_ing = doc.add_table(rows=1, cols=2)
    table_ing.style = 'Table Grid'
    hdr_cells = table_ing.rows[0].cells
    hdr_cells[0].text = 'Género'
    hdr_cells[1].text = 'Ingreso Promedio'
    
    for gen, val in ingresos.items():
        row_cells = table_ing.add_row().cells
        row_cells[0].text = str(gen)
        row_cells[1].text = f"${val:,.2f}"
        
    doc.save(os.path.join(OUTPUT_DIR, 'Reporte_Demografico_Ejecutivo.docx'))
    print("Reporte Demográfico Ejecutivo generado.")

# 3. Reporte Operativo
def generar_reporte_operativo():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: ANÁLISIS OPERATIVO TERRITORIAL', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph(
        "Se presenta el análisis de la distribución geográfica de las operaciones. Identificar los focos de mayor "
        "actividad permite optimizar la logística de servicios y la presencia comercial en las regiones clave."
    )
    
    # Actividad por Departamento
    doc.add_heading('2. Actividad por Departamento', level=1)
    deptos = df['departamento'].value_counts()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light List Accent 2'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Departamento'
    hdr_cells[1].text = 'Nivel de Operaciones'
    
    for depto, count in deptos.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(depto)
        row_cells[1].text = f"{count:,} tx"
        
    # Top Ciudades
    doc.add_heading('3. Ciudades con Mayor Volumen de Transacción', level=1)
    ciudades = df['ciudad'].value_counts().head(5)
    
    for ciudad, count in ciudades.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{ciudad}: ").bold = True
        p.add_run(f"{count:,} transacciones")
        
    doc.save(os.path.join(OUTPUT_DIR, 'Reporte_Operativo_Ejecutivo.docx'))
    print("Reporte Operativo Ejecutivo generado.")

# 4. Reporte Financiero
def generar_reporte_financiero():
    doc = Document()
    doc.add_heading('REPORTE EJECUTIVO: DESEMPEÑO FINANCIERO Y RENTABILIDAD', 0)
    
    # Resumen Ejecutivo
    doc.add_heading('1. Estados Financieros Consolidados', level=1)
    total_monto = df['monto_total'].sum()
    total_comision = df['comision'].sum()
    total_impuesto = df['impuesto'].sum()
    
    table_res = doc.add_table(rows=3, cols=2)
    table_res.style = 'Colorful List Accent 3'
    
    table_res.rows[0].cells[0].text = "Volumen Total Transado"
    table_res.rows[0].cells[1].text = f"${total_monto:,.2f}"
    
    table_res.rows[1].cells[0].text = "Total Comisiones"
    table_res.rows[1].cells[1].text = f"${total_comision:,.2f}"
    
    table_res.rows[2].cells[0].text = "Total Impuestos Recaudados"
    table_res.rows[2].cells[1].text = f"${total_impuesto:,.2f}"
    
    # Rentabilidad por Producto
    doc.add_heading('2. Carga Financiera por Categoría de Tarjeta', level=1)
    cat_tarjeta = df.groupby('categoria_tarjeta')['monto_total'].sum()
    
    table_tar = doc.add_table(rows=1, cols=2)
    table_tar.style = 'Table Grid'
    hdr_cells = table_tar.rows[0].cells
    hdr_cells[0].text = 'Tipo de Tarjeta'
    hdr_cells[1].text = 'Monto Acumulado'
    
    for cat, val in cat_tarjeta.items():
        row_cells = table_tar.add_row().cells
        row_cells[0].text = str(cat)
        row_cells[1].text = f"${val:,.2f}"
        
    doc.save(os.path.join(OUTPUT_DIR, 'Reporte_Financiero_Ejecutivo.docx'))
    print("Reporte Financiero Ejecutivo generado.")

if __name__ == "__main__":
    generar_reporte_fraudes()
    generar_reporte_demografico()
    generar_reporte_operativo()
    generar_reporte_financiero()
    print(f"\nProceso finalizado. Reportes guardados en '{OUTPUT_DIR}'")
